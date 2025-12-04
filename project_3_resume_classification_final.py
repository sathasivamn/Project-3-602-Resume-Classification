"""Project 3_Resume Classification Final.ipynb

# Imports and NLTK setup
import os
import zipfile
from pathlib import Path

import pandas as pd
import numpy as np

from docx import Document
import PyPDF2

import nltk
from nltk.corpus import stopwords
import re

from sklearn.model_selection import train_test_split
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.pipeline import Pipeline
from sklearn.linear_model import LogisticRegression
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import LinearSVC
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import (
    accuracy_score, f1_score, classification_report, confusion_matrix
)

import matplotlib.pyplot as plt
import seaborn as sns

import joblib

#Download NLTK data:
nltk.download('stopwords')
nltk.download('punkt')
stop_words = set(stopwords.words('english'))


"""

from google.colab import files

uploaded = files.upload()  # upload resume_classification.zip

ZIP_NAME = "resume_classification.zip"   # change if your zip has a different name
EXTRACT_DIR = "/content/data"
os.makedirs(EXTRACT_DIR, exist_ok=True)

if ZIP_NAME in uploaded:
    with zipfile.ZipFile(ZIP_NAME, "r") as z:
        z.extractall(EXTRACT_DIR)
    print("Zip extracted to:", EXTRACT_DIR)
else:
    raise FileNotFoundError("Upload your dataset zip using files.upload()")

# (Optional) Convert .doc → .docx with LibreOffice
base = Path(EXTRACT_DIR)

doc_files = list(base.rglob("*.doc"))
print("Found .doc files:", len(doc_files))

for doc in doc_files:
    outdir = str(doc.parent)
    cmd = f'soffice --headless --convert-to docx "{doc}" --outdir "{outdir}"'
    os.system(cmd)

print("DOC → DOCX conversion completed.")

# Define file readers (DOCX + PDF)
def read_docx(path):
    try:
        d = Document(path)
        return "\n".join(p.text for p in d.paragraphs)
    except Exception as e:
        # print("DOCX read error:", e)
        return ""

def read_pdf(path):
    try:
        r = PyPDF2.PdfReader(path)
        all_text = ""
        for p in r.pages:
            all_text += p.extract_text() or ""
        return all_text
    except Exception as e:
        # print("PDF read error:", e)
        return ""

# Build the DataFrame df
rows = []
skipped = []

all_files = [p for p in base.rglob("*") if p.is_file()]

for p in all_files:
    ext = p.suffix.lower()
    if ext not in [".pdf", ".docx"]:
        continue  # ignore other file types

    # folder name is label (e.g. 'DataScience', 'HR')
    label = p.parts[-2]

    if ext == ".docx":
        text = read_docx(str(p))
    elif ext == ".pdf":
        text = read_pdf(str(p))
    else:
        text = ""

    if text.strip():
        rows.append({
            "filename": p.name,
            "filepath": str(p),
            "label": label,
            "text": text
        })
    else:
        skipped.append(str(p))

df = pd.DataFrame(rows)

print(" DataFrame 'df' CREATED SUCCESSFULLY ")
print("Total Documents:", len(df))
print("Skipped Files:", len(skipped))
print("Columns:", df.columns.tolist())
df.head()

# EDA – Basic Exploratory Analysis
print("Columns:", df.columns.tolist())
print("\nLabel distribution:")
print(df['label'].value_counts())

plt.figure(figsize=(8,4))
df['label'].value_counts().plot(kind='bar')
plt.title("Number of resumes per class")
plt.xlabel("Class")
plt.ylabel("Count")
plt.show()

# Text length stats
df['char_count'] = df['text'].apply(len)
df['word_count_raw'] = df['text'].apply(lambda x: len(str(x).split()))

print(df[['char_count', 'word_count_raw']].describe())

plt.figure(figsize=(8,4))
df['word_count_raw'].hist(bins=30)
plt.title("Distribution of word counts in resumes")
plt.xlabel("Word count")
plt.ylabel("Frequency")
plt.show()

# Quick sample
df[['filename', 'label']].head(10)

# one full resume text sample:
sample_row = df.iloc[0]
print("Filename:", sample_row['filename'])
print("Label:", sample_row['label'])
print("--- Text snippet ---")
print(sample_row['text'][:1000])

# Text cleaning functions
# Use the same cleaning strategy for training and Streamlit prediction.
def clean_text_basic(s):
    s = str(s).lower()
    s = re.sub(r'\s+', ' ', s)            # collapse whitespace
    s = re.sub(r'http\S+', '', s)         # remove urls
    s = re.sub(r'\@\w+', '', s)           # remove @mentions
    s = re.sub(r'[^a-z0-9\s]', ' ', s)    # keep alphanumeric + space
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def remove_stopwords(s):
    tokens = s.split()
    tokens = [t for t in tokens if t not in stop_words]
    return " ".join(tokens)

df['clean_text'] = df['text'].apply(clean_text_basic)
df['clean_text_nostop'] = df['clean_text'].apply(remove_stopwords)

df['char_count_clean'] = df['clean_text'].apply(len)
df['word_count_clean'] = df['clean_text'].apply(lambda x: len(x.split()))

print(df[['char_count_clean', 'word_count_clean']].describe())

# Train/test split
X = df['clean_text']       # or 'clean_text_nostop'
y = df['label']

X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.2, random_state=42, stratify=y
)

print("Train size:", len(X_train))
print("Test size:", len(X_test))

# Train & evaluate multiple models
def train_and_eval(clf, X_train, y_train, X_test, y_test, name=None):
    pipe = Pipeline([
        ('tfidf', TfidfVectorizer(
            ngram_range=(1,2),
            max_df=0.95,
            min_df=2,
            max_features=15000
        )),
        ('clf', clf)
    ])

    pipe.fit(X_train, y_train)
    preds = pipe.predict(X_test)

    acc = accuracy_score(y_test, preds)
    f1 = f1_score(y_test, preds, average='macro')

    print(f"\n--- {name or clf.__class__.__name__} ---")
    print("Accuracy :", acc)
    print("Macro F1 :", f1)
    print(classification_report(y_test, preds))

    return pipe, acc, f1, preds

models = {
    'LogisticRegression': LogisticRegression(max_iter=2000),
    'MultinomialNB': MultinomialNB(),
    'LinearSVC': LinearSVC(max_iter=20000),
    'RandomForest': RandomForestClassifier(n_estimators=200, random_state=42)
}

results = []
pipelines = {}

for name, clf in models.items():
    pipe, acc, f1, preds = train_and_eval(clf, X_train, y_train, X_test, y_test, name=name)
    results.append({'model': name, 'accuracy': acc, 'f1_macro': f1})
    pipelines[name] = pipe

res_df = pd.DataFrame(results).sort_values('f1_macro', ascending=False).reset_index(drop=True)
res_df

# Plot comparison + confusion matrix for best model
plt.figure(figsize=(8,4))
plt.bar(res_df['model'], res_df['accuracy'])
plt.title('Model accuracy comparison')
plt.ylim(0,1)
plt.show()

plt.figure(figsize=(8,4))
plt.bar(res_df['model'], res_df['f1_macro'])
plt.title('Model macro F1 comparison')
plt.ylim(0,1)
plt.show()

# Confusion matrix:
best_model_name = res_df.loc[0, 'model']
best_pipe = pipelines[best_model_name]

y_pred = best_pipe.predict(X_test)
cm = confusion_matrix(y_test, y_pred, labels=best_pipe.classes_)

fig, ax = plt.subplots(figsize=(8,6))
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues',
            xticklabels=best_pipe.classes_,
            yticklabels=best_pipe.classes_)
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title(f'Confusion Matrix - {best_model_name}')
plt.show()

# Save the best model for deployment
pickle_filename = 'best_model.pkl'
joblib.dump(best_pipe, pickle_filename)
print(f"The best model ('{best_model_name}') has been saved to '{pickle_filename}'")

# Streamlit App (app.py)

import streamlit as st
import joblib
import io
import re

from docx import Document
import PyPDF2

import nltk
from nltk.corpus import stopwords

# Ensure NLTK stopwords are available
try:
    stop_words = set(stopwords.words('english'))
except LookupError:
    nltk.download('stopwords')
    stop_words = set(stopwords.words('english'))

# ---------------------------
# Text cleaning (same as training)
# ---------------------------

def clean_text_basic(s):
    s = str(s).lower()
    s = re.sub(r'\s+', ' ', s)            # collapse whitespace
    s = re.sub(r'http\S+', '', s)         # remove urls
    s = re.sub(r'\@\w+', '', s)           # remove @mentions
    s = re.sub(r'[^a-z0-9\s]', ' ', s)    # keep alphanumeric + space
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def remove_stopwords(s):
    tokens = s.split()
    tokens = [t for t in tokens if t not in stop_words]
    return " ".join(tokens)

def preprocess_text(s):
    s = clean_text_basic(s)
    s = remove_stopwords(s)
    return s

# ---------------------------
# File readers for uploaded resumes
# ---------------------------

def read_docx_file(file_bytes):
    file_stream = io.BytesIO(file_bytes)
    doc = Document(file_stream)
    return "\n".join(p.text for p in doc.paragraphs)

def read_pdf_file(file_bytes):
    file_stream = io.BytesIO(file_bytes)
    reader = PyPDF2.PdfReader(file_stream)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text

def read_txt_file(file_bytes):
    return file_bytes.decode("utf-8", errors="ignore")

def extract_text_from_upload(uploaded_file):
    if uploaded_file is None:
        return ""

    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".docx"):
        return read_docx_file(file_bytes)
    elif name.endswith(".pdf"):
        return read_pdf_file(file_bytes)
    elif name.endswith(".txt"):
        return read_txt_file(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please upload .pdf, .docx, or .txt")

# ---------------------------
# Load model
# ---------------------------

@st.cache_resource
def load_model():
    model = joblib.load("best_model.pkl")
    return model

model = load_model()

# ---------------------------
# Streamlit UI
# ---------------------------

st.title("📄 Resume Classification App")
st.write(
    "Upload a resume (.pdf / .docx / .txt) and the model will predict the job category."
)

uploaded_file = st.file_uploader(
    "Upload your resume file",
    type=["pdf", "docx", "txt"]
)

if uploaded_file is not None:
    st.write("**File uploaded:**", uploaded_file.name)

    if st.button("Predict Category"):
        try:
            raw_text = extract_text_from_upload(uploaded_file)

            if not raw_text.strip():
                st.error("Could not extract any text from the file.")
            else:
                st.subheader("Extracted Text (first 1000 chars)")
                st.text(raw_text[:1000])

                processed_text = preprocess_text(raw_text)
                pred_label = model.predict([processed_text])[0]
                st.subheader("Predicted Category")
                st.success(pred_label)

                # Show probabilities if available
                clf = model.named_steps.get('clf', None)
                if clf is not None and hasattr(clf, "predict_proba"):
                    probs = clf.predict_proba(
                        model.named_steps['tfidf'].transform([processed_text])
                    )[0]
                    classes = clf.classes_

                    st.subheader("Prediction Probabilities")
                    prob_table = {cls: float(p) for cls, p in zip(classes, probs)}
                    st.write(prob_table)

        except Exception as e:
            st.error(f"Error processing file: {e}")

