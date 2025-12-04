# -*- coding: utf-8 -*-

# Install Python libraries
!pip install python-docx
!pip install PyPDF2
!pip install textract
!pip install joblib scikit-learn matplotlib pandas

# Required for textract to read .doc files
!apt-get update -y
!apt-get install -y antiword unrtf poppler-utils tesseract-ocr libreoffice

from google.colab import drive
drive.mount('/content/drive')

# Example path (change to your zip folder)
# DATA_PATH = "/content/drive/MyDrive/your_folder/"

from google.colab import files
uploaded = files.upload()

import zipfile, os

ZIP_NAME = "Data set.zip"   # CHANGE if different

with zipfile.ZipFile(ZIP_NAME, 'r') as zip_ref:
    zip_ref.extractall("data")

DATA_PATH = "/content/data"

!apt-get update -y
!apt-get install -y libreoffice

import os

folder = "/content/data"  # <-- your resumes folder

for file in os.listdir(folder):
    if file.endswith(".doc"):
        input_path = os.path.join(folder, file)
        output_path = os.path.join(folder, file.replace(".doc", ".docx"))

        !soffice --headless --convert-to docx "{input_path}" --outdir "{folder}"

print("DOC → DOCX Conversion Completed!")

from docx import Document
import PyPDF2
import os

def read_docx(path):
    try:
        doc = Document(path)
        return "\n".join([p.text for p in doc.paragraphs])
    except:
        return ""

def read_pdf(path):
    try:
        reader = PyPDF2.PdfReader(path)
        txt = ''
        for page in reader.pages:
            txt += page.extract_text() or ""
        return txt
    except:
        return ""

# Optional: run this if you don't already have nltk stopwords/tokenizers
# (Run this in a terminal / notebook cell)
!pip install nltk
import nltk
nltk.download('stopwords')
nltk.download('punkt')

# MASTER CELL — CREATES df (Google Colab Safe)


import os
import zipfile
import pandas as pd
from pathlib import Path
import PyPDF2
from docx import Document

# 1. Set dataset zip and extract folder

ZIP_NAME = "Data set.zip"   # CHANGE if different
EXTRACT_DIR = "/content/data"
os.makedirs(EXTRACT_DIR, exist_ok=True)

# 2. Unzip dataset

if os.path.exists(ZIP_NAME):
    with zipfile.ZipFile(ZIP_NAME, "r") as z:
        z.extractall(EXTRACT_DIR)
else:
    raise FileNotFoundError("Upload your dataset zip using files.upload()")

# 3. Convert .doc → .docx using LibreOffice

# Install if not already installed
!apt-get update -y
!apt-get install -y libreoffice

base = Path(EXTRACT_DIR)

doc_files = list(base.rglob("*.doc"))
print("Found .doc files:", len(doc_files))

for doc in doc_files:
    outdir = str(doc.parent)
    cmd = f'soffice --headless --convert-to docx "{doc}" --outdir "{outdir}"'
    os.system(cmd)

# 4. Define readers
# ----------------------------------------------------------
def read_docx(path):
    try:
        d = Document(path)
        return "\n".join([p.text for p in d.paragraphs])
    except:
        return ""

def read_pdf(path):
    try:
        r = PyPDF2.PdfReader(path)
        all_text = ""
        for p in r.pages:
            all_text += p.extract_text() or ""
        return all_text
    except:
        return ""

# 5. Read all files and create df
# ----------------------------------------------------------
rows = []
skipped = []

all_files = [p for p in base.rglob("*") if p.is_file()]

for p in all_files:
    ext = p.suffix.lower()
    label = p.parts[-2]     # folder name = class label
    text = ""

    if ext == ".docx":
        text = read_docx(str(p))
    elif ext == ".pdf":
        text = read_pdf(str(p))
    else:
        text = ""

    if text.strip():
        rows.append({
            "filename": p.name,
            "filepath": str(p),
            "label": label,
            "text": text
        })
    else:
        skipped.append(str(p))

# Convert to df
df = pd.DataFrame(rows)


print(" DataFrame 'df' CREATED SUCCESSFULLY ")

print("Total Documents:", len(df))
print("Skipped Files:", len(skipped))
print("Columns:", df.columns.tolist())
print("Sample:")
df.head()

# Basic sanity checks and distribution
print("Columns:", df.columns.tolist())
print("Total documents:", len(df))
print("Sample labels and counts:")
print(df['label'].value_counts())

# Show a few sample filenames & labels
display(df[['filename','label']].head(10))

# Clean text and create features useful for EDA and modeling
import re
from nltk.corpus import stopwords
stop_words = set(stopwords.words('english'))

def clean_text_basic(s):
    s = str(s).lower()
    s = re.sub(r'\s+', ' ', s)            # collapse whitespace
    s = re.sub(r'http\S+', '', s)         # remove urls
    s = re.sub(r'\@\w+', '', s)           # remove @mentions
    s = re.sub(r'[^a-z0-9\s]', ' ', s)    # keep alphanumeric
    s = re.sub(r'\s+', ' ', s).strip()
    return s

def remove_stopwords(s):
    tokens = s.split()
    tokens = [t for t in tokens if t not in stop_words]
    return " ".join(tokens)

# Apply cleaning (keep both raw_clean and optionally stopword-removed)
df['clean_text'] = df['text'].apply(clean_text_basic)
df['clean_text_nostop'] = df['clean_text'].apply(remove_stopwords)

# Show stats
df['char_count'] = df['clean_text'].apply(len)
df['word_count'] = df['clean_text'].apply(lambda x: len(x.split()))
print(df[['char_count','word_count']].describe())

# Create train/test split
from sklearn.model_selection import train_test_split

X = df['clean_text']            # you can also try 'clean_text_nostop'
y = df['label']

# stratify to preserve class proportions
X_train, X_test, y_train, y_test = train_test_split(
    X, y, test_size=0.20, random_state=42, stratify=y
)

print("Train:", len(X_train), "Test:", len(X_test))

# We'll train multiple models with the same TF-IDF pipeline
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.pipeline import Pipeline
from sklearn.linear_model import LogisticRegression
from sklearn.naive_bayes import MultinomialNB
from sklearn.svm import LinearSVC
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import accuracy_score, f1_score, classification_report, confusion_matrix

def train_and_eval(clf, X_train, y_train, X_test, y_test, name=None):
    pipe = Pipeline([
        ('tfidf', TfidfVectorizer(ngram_range=(1,2), max_df=0.95, min_df=2, max_features=15000)),
        ('clf', clf)
    ])
    pipe.fit(X_train, y_train)
    preds = pipe.predict(X_test)
    acc = accuracy_score(y_test, preds)
    f1 = f1_score(y_test, preds, average='macro')
    print(f"--- {name or clf.__class__.__name__} ---")
    print("Accuracy:", acc)
    print("Macro F1:", f1)
    print(classification_report(y_test, preds))
    return pipe, acc, f1, preds

# Create models to evaluate
models = {
    'LogisticRegression': LogisticRegression(max_iter=2000),
    'MultinomialNB': MultinomialNB(),
    'LinearSVC': LinearSVC(max_iter=20000),
    'RandomForest': RandomForestClassifier(n_estimators=200, random_state=42)
}

# Train all models and collect results
results = []
pipelines = {}

for name, clf in models.items():
    pipe, acc, f1, preds = train_and_eval(clf, X_train, y_train, X_test, y_test, name=name)
    results.append({'model': name, 'accuracy': acc, 'f1_macro': f1})
    pipelines[name] = pipe

# Create comparison DataFrame
import pandas as pd
res_df = pd.DataFrame(results).sort_values('f1_macro', ascending=False).reset_index(drop=True)
res_df

# Commented out IPython magic to ensure Python compatibility.
import matplotlib.pyplot as plt
# %matplotlib inline

# Bar plots
plt.figure(figsize=(8,4))
plt.bar(res_df['model'], res_df['accuracy'])
plt.title('Model accuracy comparison')
plt.ylim(0,1)
plt.show()

plt.figure(figsize=(8,4))
plt.bar(res_df['model'], res_df['f1_macro'])
plt.title('Model macro F1 comparison')
plt.ylim(0,1)
plt.show()

# Confusion matrix for the best model
best_model = res_df.loc[0, 'model']
best_pipe = pipelines[best_model]
y_pred = best_pipe.predict(X_test)

cm = confusion_matrix(y_test, y_pred, labels=best_pipe.classes_)
fig, ax = plt.subplots(figsize=(8,6))
import seaborn as sns
sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=best_pipe.classes_, yticklabels=best_pipe.classes_)
plt.xlabel('Predicted')
plt.ylabel('True')
plt.title(f'Confusion Matrix - {best_model}')
plt.show()

import joblib

# Get the best performing model pipeline
best_model_name = res_df.loc[0, 'model']
best_pipe = pipelines[best_model_name]

# Define the filename for the pickle file
pickle_filename = 'best_model.pkl'

# Dump the model to the pickle file
joblib.dump(best_pipe, pickle_filename)

print(f"The best model ('{best_model_name}') has been saved to '{pickle_filename}'")

import joblib
!pip install joblib
!pip install streamlit

!streamlit run app.py

import docx
print("python-docx installed successfully!")

